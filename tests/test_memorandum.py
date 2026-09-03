"""Het reviewmemorandum als document.

De bevindingen worden hier zelf opgebouwd in plaats van uit een volledige
analyse gehaald: dan staat in de test wat er in het document hoort te komen, en
gaat zij niet stuk zodra een controle een bevinding meer of minder oplevert.
Alle bestanden zijn synthetisch.
"""
from __future__ import annotations

from datetime import date
from io import BytesIO

import pandas as pd

from auditfile.demo import demopaar
from auditfile.findings import (
    AFGEHANDELDE_STATUSSEN,
    BEVINDING_COLUMNS,
    REVIEWSTATUSSEN,
    SIGNAAL,
    TE_BEOORDELEN,
    Bevinding,
    Materialiteit,
    grondslag_omzet,
    naar_frame,
    pas_review_toe,
    verzamel_bevindingen,
)
from auditfile.integrity import KRITIEK, NIET_MOGELIJK, WAARSCHUWING
from auditfile.memorandum import (
    DOCX_GENUMMERD,
    DOCX_LIJST,
    DOCX_ONDERTITEL,
    ONDER_DREMPEL,
    SAMENVATTING_MAX,
    bouw_memorandum,
    memorandum_markdown,
    memorandumnaam,
    naar_docx,
    naar_markdown,
    op_gewicht,
)
from auditfile.parsing import parse_auditfile

MATERIALITEIT = Materialiteit(absoluut=1000.0, relatief_pct=1.0, grondslag=500_000.0)


def _demopaar():
    vorig_bytes, huidig_bytes = demopaar()
    return (
        parse_auditfile("demo_vorig_jaar.xaf", vorig_bytes),
        parse_auditfile("demo_huidig_jaar.xaf", huidig_bytes),
    )


def _frame(*bevindingen: Bevinding, review: dict | None = None) -> pd.DataFrame:
    frame = naar_frame(list(bevindingen), MATERIALITEIT)
    return pas_review_toe(frame, review)


def _sectiekoppen(memo) -> list[str]:
    return [sectie.kop for sectie in memo.secties]


def _sectie(memo, fragment: str):
    for sectie in memo.secties:
        if fragment.lower() in sectie.kop.lower():
            return sectie
    raise AssertionError(f"Geen sectie met '{fragment}' in {_sectiekoppen(memo)}")


# --- Volledigheid -----------------------------------------------------------


def test_elke_bevinding_komt_in_het_document(af_40):
    """Wat in de lijst staat, mag niet stil uit het stuk verdwijnen."""
    frame = _frame(
        Bevinding("Btw", "Rondrekening sluit niet aan", KRITIEK, bedrag=12_000.0),
        Bevinding("Fiscaal", "Rekening-courant met de dga", WAARSCHUWING, bedrag=400_000.0),
        Bevinding("Boekingen", "Rond bedrag", SIGNAAL, bedrag=25_000.0),
        Bevinding("Bestandsgegevens", "Openstaande posten niet te bepalen", NIET_MOGELIJK),
        review={},
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)
    tekst = naar_markdown(memo)

    assert len(memo.punten) == len(frame)
    for onderwerp in frame["onderwerp"]:
        assert onderwerp in tekst


def test_de_volledige_analyse_levert_een_document_op():
    """De echte verzamelaar, om te zien dat elk soort bevinding erdoor komt."""
    vorig, huidig = _demopaar()
    materialiteit = Materialiteit(grondslag=grondslag_omzet(huidig))
    frame = pas_review_toe(verzamel_bevindingen(huidig, vorig, materialiteit=materialiteit))

    memo = bouw_memorandum(huidig, frame, materialiteit, vorig=vorig)
    tekst = naar_markdown(memo)

    assert len(memo.punten) == len(frame)
    assert huidig.boekjaar in tekst
    assert vorig.bestandsnaam in tekst
    # Geen enkel punt mag zonder onderwerp in het stuk staan.
    assert all(punt.onderwerp for punt in memo.punten)


def test_zonder_bevindingen_blijft_het_document_geldig(af_40):
    leeg = pas_review_toe(pd.DataFrame(columns=BEVINDING_COLUMNS))

    tekst = memorandum_markdown(af_40, leeg, MATERIALITEIT)

    assert "Reviewmemorandum" in tekst
    assert "geen aandachtspunten" in tekst
    # Ook zonder bevindingen hoort de verantwoording er te staan.
    assert "## Verantwoording" in tekst


# --- Wat niet kon worden vastgesteld ----------------------------------------


def test_niet_mogelijk_krijgt_een_eigen_sectie(af_40):
    frame = _frame(
        Bevinding("Btw", "Aansluiting met de aangifte niet gemaakt", NIET_MOGELIJK),
        Bevinding("Boekingen", "Rond bedrag", SIGNAAL, bedrag=25_000.0),
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)
    sectie = _sectie(memo, "niet kon worden vastgesteld")

    assert [punt.onderwerp for punt in sectie.punten] == [
        "Aansluiting met de aangifte niet gemaakt"
    ]
    assert "ontbreekt ook het bewijs" in sectie.inleiding


def test_niet_mogelijk_blijft_zichtbaar_ook_na_beoordeling(af_40):
    """Een afgehandelde status mag een onmogelijke controle niet wegmoffelen."""
    bevinding = Bevinding("Btw", "Aansluiting met de aangifte niet gemaakt", NIET_MOGELIJK)
    frame = _frame(
        bevinding,
        review={bevinding.sleutel: {"status": "Niet van toepassing", "notitie": "Geen btw."}},
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)
    sectie = _sectie(memo, "niet kon worden vastgesteld")

    assert [punt.onderwerp for punt in sectie.punten] == [bevinding.onderwerp]
    assert "Al beoordeeld" not in _sectiekoppen(memo)
    # De vastgelegde beoordeling gaat wel mee.
    assert "Niet van toepassing" in naar_markdown(memo)


def test_de_sectie_staat_er_ook_zonder_onmogelijke_controles(af_40):
    frame = _frame(Bevinding("Boekingen", "Rond bedrag", SIGNAAL, bedrag=25_000.0))

    sectie = _sectie(bouw_memorandum(af_40, frame, MATERIALITEIT), "niet kon worden vastgesteld")

    assert not sectie.punten
    assert "kon op deze bestanden worden uitgevoerd" in sectie.inleiding


# --- Beoordeling ------------------------------------------------------------


def test_afgehandelde_bevinding_verhuist_naar_achteren(af_40):
    opgelost = Bevinding("Btw", "Rondrekening sluit niet aan", KRITIEK, bedrag=90_000.0)
    actie = Bevinding("Fiscaal", "Rekening-courant met de dga", WAARSCHUWING, bedrag=400_000.0)
    frame = _frame(
        opgelost,
        actie,
        review={
            opgelost.sleutel: {"status": "Opgelost", "notitie": "Suppletie ingediend."},
            actie.sleutel: {"status": "Actie nodig", "notitie": "Navraag bij de klant."},
        },
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)

    afgehandeld = _sectie(memo, "Al beoordeeld")
    assert [punt.onderwerp for punt in afgehandeld.punten] == [opgelost.onderwerp]
    # "Actie nodig" is geen afgehandelde status en blijft tussen de punten die
    # aandacht vragen staan.
    aandacht = _sectie(memo, "Waarschuwing")
    assert [punt.onderwerp for punt in aandacht.punten] == [actie.onderwerp]
    assert "Suppletie ingediend." in naar_markdown(memo)
    assert "Navraag bij de klant." in naar_markdown(memo)


def test_de_samenvatting_gaat_over_wat_nog_aandacht_vraagt(af_40):
    """Een afgehandeld punt met een groot bedrag hoort niet bovenaan."""
    opgelost = Bevinding("Btw", "Al opgelost", KRITIEK, bedrag=900_000.0)
    open_punt = Bevinding("Boekingen", "Nog te beoordelen", SIGNAAL, bedrag=5_000.0)
    frame = _frame(
        opgelost,
        open_punt,
        review={opgelost.sleutel: {"status": "Beoordeeld, geen actie", "notitie": ""}},
    )
    samenvatting = _sectie(bouw_memorandum(af_40, frame, MATERIALITEIT), "Samenvatting")

    assert any("Nog te beoordelen" in regel for regel in samenvatting.regels)
    assert not any("Al opgelost" in regel for regel in samenvatting.regels)
    assert "1 bevinding is al beoordeeld" in samenvatting.inleiding


def test_zonder_vastgelegde_beoordeling_staat_er_geen_status(af_40):
    frame = _frame(Bevinding("Boekingen", "Rond bedrag", SIGNAAL, bedrag=25_000.0))

    tekst = memorandum_markdown(af_40, frame, MATERIALITEIT)

    assert f"Beoordeling: {TE_BEOORDELEN}" not in tekst
    assert "hebben nog geen vastgelegde beoordeling" in tekst


def test_afgehandelde_statussen_bestaan_echt():
    """Zonder deze controle zou een hernoemde status stil niets meer doen."""
    assert set(AFGEHANDELDE_STATUSSEN) < set(REVIEWSTATUSSEN)
    assert TE_BEOORDELEN not in AFGEHANDELDE_STATUSSEN
    assert "Actie nodig" not in AFGEHANDELDE_STATUSSEN


# --- Ordening en materialiteit ----------------------------------------------


def test_zwaarste_eerst_en_zonder_bedrag_niet_achteraan():
    frame = _frame(
        Bevinding("Boekingen", "Klein signaal", SIGNAAL, bedrag=10.0),
        Bevinding("Boekingen", "Signaal zonder bedrag", SIGNAAL),
        Bevinding("Boekingen", "Groot signaal", SIGNAAL, bedrag=50_000.0),
        Bevinding("Btw", "Kritiek punt", KRITIEK, bedrag=2_000.0),
    )
    volgorde = list(op_gewicht(frame)["onderwerp"])

    assert volgorde == [
        "Kritiek punt",
        "Groot signaal",
        "Signaal zonder bedrag",
        "Klein signaal",
    ]


def test_onder_de_drempel_wordt_gemarkeerd_en_niet_weggelaten(af_40):
    klein = Bevinding("Boekingen", "Kleine post", SIGNAAL, bedrag=10.0)
    frame = _frame(klein, Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0))

    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)
    punt = next(punt for punt in memo.punten if punt.onderwerp == "Kleine post")

    assert not punt.boven_drempel
    assert ONDER_DREMPEL in punt.aanduiding
    assert "Kleine post" in naar_markdown(memo)


def test_de_drempel_staat_met_zijn_opbouw_in_het_stuk(af_40):
    frame = _frame(Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0))

    tekst = memorandum_markdown(af_40, frame, MATERIALITEIT)

    # De hoogste van de twee grenzen geldt: 1% van 500.000 is 5.000.
    assert "€ 5.000,00" in tekst
    assert "€ 1.000,00" in tekst
    assert "1% van de grondslag € 500.000,00" in tekst
    assert "geen norm uit wet of standaard" in tekst


# --- Nummering en samenvatting ----------------------------------------------


def test_de_nummering_loopt_door_en_is_uniek(af_40):
    frame = _frame(
        Bevinding("Btw", "Kritiek punt", KRITIEK, bedrag=50_000.0),
        Bevinding("Boekingen", "Signaal", SIGNAAL, bedrag=25_000.0),
        Bevinding("Bestandsgegevens", "Niet te bepalen", NIET_MOGELIJK),
    )
    nummers = [punt.nummer for punt in bouw_memorandum(af_40, frame, MATERIALITEIT).punten]

    assert nummers == [1, 2, 3]


def test_de_samenvatting_noemt_niet_meer_dan_een_handvol(af_40):
    frame = _frame(
        *[
            Bevinding("Boekingen", f"Signaal {index}", SIGNAAL, bedrag=100_000.0 - index)
            for index in range(SAMENVATTING_MAX + 3)
        ]
    )
    samenvatting = _sectie(bouw_memorandum(af_40, frame, MATERIALITEIT), "Samenvatting")

    assert len(samenvatting.regels) == SAMENVATTING_MAX
    assert samenvatting.regels[0].startswith("Signaal 0")
    assert "De overige 3 punten" in samenvatting.slot


# --- Kop en verantwoording --------------------------------------------------


def test_de_kop_benoemt_de_bestanden_en_de_versies(af_40, af_32):
    frame = _frame(Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0))

    memo = bouw_memorandum(
        af_40, frame, MATERIALITEIT, vorig=af_32, opsteller="A. Beoordelaar",
        opgesteld_op=date(2026, 9, 3),
    )
    kenmerken = dict(_sectie(memo, "Dossier").kenmerken)

    assert af_40.bestandsnaam in kenmerken["Auditfile huidig jaar"]
    assert "XAF 4.0" in kenmerken["Auditfile huidig jaar"]
    assert "XAF 3.2" in kenmerken["Auditfile vorig jaar"]
    assert kenmerken["Opgesteld op"] == "03-09-2026"
    assert kenmerken["Opgesteld door"] == "A. Beoordelaar"


def test_een_ontbrekend_vergelijkingsbestand_wordt_benoemd(af_40):
    frame = _frame(Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0))

    kenmerken = dict(_sectie(bouw_memorandum(af_40, frame, MATERIALITEIT), "Dossier").kenmerken)

    assert "niet geladen" in kenmerken["Auditfile vorig jaar"]


def test_de_verantwoording_noemt_bewijsniveau_en_rgs_dekking(af_40):
    frame = _frame(Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0))

    regels = " ".join(_sectie(bouw_memorandum(af_40, frame, MATERIALITEIT), "Verantwoording").regels)

    assert "Openstaande posten zijn te bepalen op niveau" in regels
    assert "RGS-code" in regels
    assert "geen accountantsverklaring" in regels


def test_de_bestandsnaam_bevat_geen_klantnaam(af_40):
    naam = memorandumnaam(af_40)

    assert naam.endswith(".md")
    assert af_40.bedrijfsnaam
    assert af_40.bedrijfsnaam.split()[0].lower() not in naam.lower()


# --- Word-uitvoer -----------------------------------------------------------
#
# De tweede renderer op dezelfde ``Memorandum``. Het bestand wordt hier weer
# ingelezen zoals Word het zou openen: een test op de bytes zou een geldig
# document niet van een leeg document onderscheiden.


def _word(inhoud: bytes):
    from docx import Document

    return Document(BytesIO(inhoud))


def _alineas(inhoud: bytes) -> list[tuple[str, str]]:
    """Elke alinea uit het Word-bestand als stijl met tekst."""
    return [(alinea.style.name, alinea.text) for alinea in _word(inhoud).paragraphs]


def _woordtekst(inhoud: bytes) -> str:
    return "\n".join(tekst for _, tekst in _alineas(inhoud))


def test_elk_onderwerp_komt_ook_in_het_word_document(af_40):
    """Wat in de tekst staat, mag in Word niet stil wegvallen."""
    beoordeeld = Bevinding("Btw", "Rondrekening sluit niet aan", KRITIEK, bedrag=90_000.0)
    bevindingen = [
        beoordeeld,
        Bevinding("Fiscaal", "Rekening-courant met de dga", WAARSCHUWING, bedrag=400_000.0),
        Bevinding("Boekingen", "Rond bedrag", SIGNAAL, bedrag=25_000.0, aantal_regels=16),
        Bevinding("Boekingen", "Kleine post", SIGNAAL, bedrag=40.0),
        Bevinding("Bestandsgegevens", "Openstaande posten niet te bepalen", NIET_MOGELIJK),
    ]
    frame = _frame(
        *bevindingen,
        review={beoordeeld.sleutel: {"status": "Opgelost", "notitie": "Suppletie ingediend."}},
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)

    tekst = _woordtekst(naar_docx(memo))

    for bevinding in bevindingen:
        assert bevinding.onderwerp in tekst
    # De beoordeling en de markering onder de drempel horen er ook in te staan.
    assert "Suppletie ingediend." in tekst
    assert ONDER_DREMPEL in tekst


def test_de_volledige_analyse_levert_een_word_document_op():
    """Op de echte bevindingenlijst en niet op een handvol zelfgemaakte punten."""
    vorig, huidig = _demopaar()
    materialiteit = Materialiteit(grondslag=grondslag_omzet(huidig))
    frame = pas_review_toe(verzamel_bevindingen(huidig, vorig, materialiteit=materialiteit))

    memo = bouw_memorandum(huidig, frame, materialiteit, vorig=vorig)
    tekst = _woordtekst(naar_docx(memo))

    assert not frame.empty
    for onderwerp in frame["onderwerp"]:
        assert onderwerp in tekst
    assert vorig.bestandsnaam in tekst


def test_het_word_document_volgt_de_koppen_van_het_memorandum(af_40):
    frame = _frame(
        Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0),
        Bevinding("Bestandsgegevens", "Niet te bepalen", NIET_MOGELIJK),
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)

    alineas = _alineas(naar_docx(memo))

    assert ("Title", memo.titel) in alineas
    assert (DOCX_ONDERTITEL, memo.ondertitel) in alineas
    assert [tekst for stijl, tekst in alineas if stijl == "Heading 1"] == [
        sectie.kop for sectie in memo.secties
    ]
    # Elk punt staat als kop, met hetzelfde doorlopende nummer als in de tekst.
    assert [tekst for stijl, tekst in alineas if stijl == "Heading 3"] == [
        f"{punt.nummer}. {punt.aanduiding}" for punt in memo.punten
    ]


def test_de_kenmerken_staan_met_een_vet_label(af_40):
    frame = _frame(Bevinding("Btw", "Grote post", KRITIEK, bedrag=50_000.0))
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT, opsteller="A. Beoordelaar")

    kenmerken = {
        alinea.runs[0].text.rstrip(": "): (bool(alinea.runs[0].bold), alinea.runs[1].text)
        for alinea in _word(naar_docx(memo)).paragraphs
        if alinea.style.name == DOCX_LIJST and len(alinea.runs) == 2
    }

    assert kenmerken["Opgesteld door"] == (True, "A. Beoordelaar")
    assert kenmerken["Boekjaar"] == (True, af_40.boekjaar)


def test_de_herkomstregel_is_cursief(af_40):
    frame = _frame(
        Bevinding(
            "Btw",
            "Grote post",
            KRITIEK,
            bedrag=50_000.0,
            methode="de RGS-code",
            pagina="Btw-analyse",
        )
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)
    punt = memo.punten[0]
    assert punt.herkomst

    cursief = [
        alinea.text
        for alinea in _word(naar_docx(memo)).paragraphs
        if alinea.runs and all(run.italic for run in alinea.runs)
    ]

    assert cursief == [punt.herkomst]


def test_de_samenvatting_houdt_in_word_de_nummers_van_de_punten(af_40):
    """De nummers staan in de tekst, niet in de nummering van Word.

    Zij verwijzen naar de punten verderop; automatische nummering zou dat
    verband verschuiven zodra iemand in het document een regel weghaalt.
    """
    frame = _frame(
        Bevinding("Btw", "Kritiek punt", KRITIEK, bedrag=50_000.0),
        Bevinding("Boekingen", "Signaal", SIGNAAL, bedrag=25_000.0),
    )
    memo = bouw_memorandum(af_40, frame, MATERIALITEIT)

    genummerd = [tekst for stijl, tekst in _alineas(naar_docx(memo)) if stijl == DOCX_GENUMMERD]

    assert genummerd == [f"{punt.nummer}. {punt.aanduiding}" for punt in memo.punten]


def test_zonder_bevindingen_blijft_het_word_document_geldig(af_40):
    leeg = pas_review_toe(pd.DataFrame(columns=BEVINDING_COLUMNS))
    memo = bouw_memorandum(af_40, leeg, MATERIALITEIT)

    alineas = _alineas(naar_docx(memo))

    assert ("Title", memo.titel) in alineas
    assert "geen aandachtspunten" in _woordtekst(naar_docx(memo))
    assert ("Heading 1", "Verantwoording") in alineas


def test_de_bestandsnaam_kan_ook_een_docx_zijn(af_40):
    assert memorandumnaam(af_40, "docx").endswith(".docx")
    # Met of zonder punt ervoor, en nog steeds zonder klantnaam.
    assert memorandumnaam(af_40, ".docx") == memorandumnaam(af_40, "docx")
    assert af_40.bedrijfsnaam.split()[0].lower() not in memorandumnaam(af_40, "docx").lower()
