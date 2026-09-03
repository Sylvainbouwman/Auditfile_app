"""Het reviewmemorandum als document.

``findings.py`` levert alle bevindingen in één vorm; die tabel is bruikbaar om
in te werken maar niet om te lezen. Een memorandum is iets anders dan een
gefilterde tabel: het zegt waar de bestanden vandaan komen, onder welke
materialiteit is gekeken, wat de zwaarste punten zijn en, uitdrukkelijk, wat
niet kon worden vastgesteld.

Twee lagen
----------
``bouw_memorandum()`` maakt uit de bevindingen een ``Memorandum``: secties met
punten en regels, en geen opmaak. ``naar_markdown()`` en ``naar_docx()`` zetten
dat om naar tekst en naar een Word-bestand. Alle formulering en ordening zit
daarmee in de eerste laag en is als tekst te testen; een renderer erbij geeft
hetzelfde stuk een andere vorm en nooit een tweede versie van dezelfde zinnen.
Staat een zin in een renderer, dan staat zij op de verkeerde plek.

Ordening
--------
Het memorandum sorteert zelf: eerst ernst, dan boven de drempel vóór eronder,
dan bedrag aflopend. ``naar_frame()`` sorteert op ernst en bedrag, waardoor een
bevinding zonder bedrag onderaan haar ernstgroep zakt terwijl zij juist altijd
meetelt. In een tabel is dat te overzien, in een doorlopend document niet.

Wat waar terechtkomt
--------------------
Een bevinding met ernst ``niet mogelijk`` staat altijd in de sectie over wat
niet kon worden vastgesteld, ook wanneer er een beoordeling aan hangt: een
memorandum dat zwijgt over een controle die niet kon worden uitgevoerd, wekt de
indruk dat er niets aan de hand is. Een bevinding die de gebruiker zelf heeft
afgehandeld verhuist naar een eigen sectie achteraan, met haar status en
notitie. Niets valt weg, ook niet onder de materialiteitsdrempel; wat onder de
drempel valt wordt gemarkeerd.

Signalen, geen oordelen: de punten benoemen wat er is gezien en wat beoordeling
vraagt. De fiscale conclusie is aan de beoordelaar.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from io import BytesIO

import pandas as pd

from .findings import (
    AFGEHANDELDE_STATUSSEN,
    ERNST_ORDE,
    SIGNAAL,
    TE_BEOORDELEN,
    Materialiteit,
    pas_review_toe,
)
from .integrity import KRITIEK, NIET_MOGELIJK, WAARSCHUWING
from .model import Auditfile
from .notatie import datum_nl, euro

# Hoeveel punten de samenvatting vooraan noemt. Meer dan dit is geen
# samenvatting meer; de volledige lijst staat er direct onder.
SAMENVATTING_MAX = 7

# De kop boven elke ernstgroep, met de betekenis erbij. Wie het memorandum
# leest heeft de app er niet naast en kent de vier niveaus niet.
ERNST_KOP: dict[str, str] = {
    KRITIEK: "Kritiek: zonder oplossing zijn de cijfers niet te gebruiken",
    WAARSCHUWING: "Waarschuwing: afwijking die beoordeling vraagt",
    SIGNAAL: "Signaal: iets om naar te kijken",
}

ONDER_DREMPEL = "onder de drempel"


def _tekst(waarde) -> str:
    if waarde is None or (isinstance(waarde, float) and pd.isna(waarde)):
        return ""
    return str(waarde).strip()


def _aantal(getal: int, enkelvoud: str, meervoud: str) -> str:
    """Een telling met het juiste woord erachter.

    Een document leest niemand met "1 bevinding(en)" erin; die vorm hoort bij
    een logregel en niet bij een stuk dat in een dossier komt.
    """
    return f"{getal} {enkelvoud if getal == 1 else meervoud}"


@dataclass(frozen=True)
class Punt:
    """Eén bevinding zoals zij in het document staat.

    Een eigen vorm naast ``Bevinding``, omdat hier het nummer bij hoort waarmee
    het punt in het document te vinden is en de beoordeling die de gebruiker
    eraan heeft gehangen.
    """

    nummer: int
    onderwerp: str
    ernst: str
    categorie: str = ""
    bedrag: float | None = None
    rekening: str = ""
    aantal_regels: int | None = None
    boven_drempel: bool = True
    methode: str = ""
    toelichting: str = ""
    pagina: str = ""
    status: str = TE_BEOORDELEN
    notitie: str = ""

    @property
    def aanduiding(self) -> str:
        """Het punt in één regel: onderwerp, bedrag, rekening, weging."""
        tussen = []
        if self.bedrag is not None and not pd.isna(self.bedrag):
            tussen.append(euro(self.bedrag))
        if self.rekening:
            tussen.append(f"rekening {self.rekening}")
        if self.aantal_regels is not None and not pd.isna(self.aantal_regels):
            tussen.append(_aantal(int(self.aantal_regels), "regel", "regels"))
        if not self.boven_drempel:
            tussen.append(ONDER_DREMPEL)
        if not tussen:
            return self.onderwerp
        return f"{self.onderwerp} ({', '.join(tussen)})"

    @property
    def herkomst(self) -> str:
        """Waar het punt vandaan komt, in een regel; leeg als dat niets zegt.

        Deze zin staat hier en niet in een renderer, want elke renderer zet
        haar neer en twee plaatsen worden twee formuleringen.
        """
        delen = []
        if self.categorie:
            delen.append(f"Onderdeel {self.categorie}")
        if self.methode:
            delen.append(f"gebaseerd op {self.methode}")
        if self.pagina:
            delen.append(f"onderbouwing op de pagina {self.pagina}")
        if not delen:
            return ""
        return f"{'; '.join(delen)}."

    @property
    def beoordeling(self) -> str:
        """De vastgelegde beoordeling, of leeg als er niets is vastgelegd."""
        if self.status == TE_BEOORDELEN and not self.notitie:
            return ""
        tekst = f"Beoordeling: {self.status}."
        if self.notitie:
            tekst = f"{tekst} Notitie: {self.notitie}"
        return tekst


@dataclass(frozen=True)
class Sectie:
    """Een kop met wat eronder hoort: een inleiding, kenmerken, regels, punten."""

    kop: str
    inleiding: str = ""
    kenmerken: tuple[tuple[str, str], ...] = ()
    regels: tuple[str, ...] = ()
    punten: tuple[Punt, ...] = ()
    genummerd: bool = False
    # Een afsluitende zin onder de opsomming. Zij staat apart omdat een
    # slotzin in een genummerde lijst een nummer zou krijgen dat naar een punt
    # lijkt te verwijzen.
    slot: str = ""


@dataclass(frozen=True)
class Memorandum:
    """Het document, nog zonder opmaak."""

    titel: str
    ondertitel: str = ""
    secties: tuple[Sectie, ...] = field(default_factory=tuple)

    @property
    def punten(self) -> tuple[Punt, ...]:
        """Alle punten, in de volgorde waarin ze in het document staan."""
        return tuple(punt for sectie in self.secties for punt in sectie.punten)


# --- Ordening ---------------------------------------------------------------


def op_gewicht(bevindingen: pd.DataFrame) -> pd.DataFrame:
    """Zwaarste eerst: ernst, dan boven de drempel, dan bedrag.

    De tweede sleutel is er omdat een bevinding zonder bedrag altijd meetelt: op
    bedrag alleen zou zij achter een klein en onbelangrijk bedrag eindigen.
    """
    if bevindingen.empty:
        return bevindingen.copy()
    orde = {ernst: index for index, ernst in enumerate(ERNST_ORDE)}
    frame = bevindingen.copy()
    frame["_ernst"] = frame["ernst"].map(lambda ernst: orde.get(ernst, len(orde)))
    # Onder de drempel sorteert achteraan, dus de vlag wordt omgekeerd.
    frame["_onder"] = ~frame["boven_drempel"].astype(bool)
    frame["_bedrag"] = pd.to_numeric(frame["bedrag"], errors="coerce").abs().fillna(0.0)
    return (
        frame.sort_values(
            ["_ernst", "_onder", "_bedrag"], ascending=[True, True, False], kind="stable"
        )
        .drop(columns=["_ernst", "_onder", "_bedrag"])
        .reset_index(drop=True)
    )


def _naar_punt(rij: pd.Series, nummer: int) -> Punt:
    bedrag = pd.to_numeric(rij.get("bedrag"), errors="coerce")
    regels = pd.to_numeric(rij.get("aantal_regels"), errors="coerce")
    return Punt(
        nummer=nummer,
        onderwerp=_tekst(rij.get("onderwerp")),
        ernst=_tekst(rij.get("ernst")),
        categorie=_tekst(rij.get("categorie")),
        bedrag=None if pd.isna(bedrag) else float(bedrag),
        rekening=_tekst(rij.get("rekening")),
        aantal_regels=None if pd.isna(regels) else int(regels),
        boven_drempel=bool(rij.get("boven_drempel", True)),
        methode=_tekst(rij.get("methode")),
        toelichting=_tekst(rij.get("toelichting")),
        pagina=_tekst(rij.get("pagina")),
        status=_tekst(rij.get("status")) or TE_BEOORDELEN,
        notitie=_tekst(rij.get("notitie")),
    )


# --- De secties -------------------------------------------------------------


def _sectie_dossier(
    huidig: Auditfile, vorig: Auditfile | None, opsteller: str, opgesteld_op: date
) -> Sectie:
    kenmerken: list[tuple[str, str]] = [
        ("Onderneming", huidig.bedrijfsnaam or "niet in het bestand vermeld"),
        ("Boekjaar", huidig.boekjaar or "niet in het bestand vermeld"),
    ]
    start = datum_nl(huidig.header.get("startDate", ""))
    eind = datum_nl(huidig.header.get("endDate", ""))
    if start and eind:
        kenmerken.append(("Periode", f"{start} tot en met {eind}"))
    kenmerken.append(
        ("Auditfile huidig jaar", f"{huidig.bestandsnaam} (XAF {huidig.xaf_versie or 'onbekend'})")
    )
    if vorig is not None:
        kenmerken.append(
            (
                "Auditfile vorig jaar",
                f"{vorig.bestandsnaam} (XAF {vorig.xaf_versie or 'onbekend'}), "
                f"boekjaar {vorig.boekjaar or 'onbekend'}",
            )
        )
    else:
        kenmerken.append(
            (
                "Auditfile vorig jaar",
                "niet geladen; de jaar-op-jaar vergelijking en de jaarovergang "
                "ontbreken daarmee in dit memorandum",
            )
        )
    if huidig.valuta:
        kenmerken.append(("Valuta", huidig.valuta))
    kenmerken.append(("Opgesteld op", datum_nl(opgesteld_op) or opgesteld_op.strftime("%d-%m-%Y")))
    if opsteller:
        kenmerken.append(("Opgesteld door", opsteller))
    return Sectie(kop="Dossier en bestanden", kenmerken=tuple(kenmerken))


def _sectie_uitgangspunten(materialiteit: Materialiteit) -> Sectie:
    relatief = abs(float(materialiteit.grondslag)) * float(materialiteit.relatief_pct) / 100.0
    regels = [
        f"De gebruikte materialiteitsdrempel is {euro(materialiteit.drempel)}: de hoogste van "
        f"{euro(materialiteit.absoluut)} en {materialiteit.relatief_pct:g}% van de grondslag "
        f"{euro(materialiteit.grondslag)}, dus {euro(relatief)}. Dit is een werkafspraak van "
        "de opsteller en geen norm uit wet of standaard.",
        f"Een bevinding onder die drempel is met “{ONDER_DREMPEL}” gemarkeerd en niet "
        "weggelaten. Een bevinding zonder bedrag is niet te wegen en telt daarom altijd mee.",
        "De analyse werkt op het auditfile en niet op de jaarrekening. Zij ziet dus wat er is "
        "geboekt, en niet wat er buiten de administratie om is vastgelegd.",
        "Wat hieronder staat zijn signalen: waarnemingen die beoordeling vragen. De fiscale en "
        "jaarrekeningtechnische conclusie is aan de beoordelaar en staat niet in dit stuk.",
    ]
    return Sectie(kop="Uitgangspunten", regels=tuple(regels))


def _sectie_samenvatting(
    punten: tuple[Punt, ...], niet_mogelijk: int, afgehandeld: int, vorig_geladen: bool
) -> Sectie:
    bron = "deze twee auditfiles" if vorig_geladen else "dit auditfile"
    if not punten:
        inleiding = (
            f"Op basis van {bron} zijn geen aandachtspunten benoemd die nog beoordeling "
            "vragen. Dat is zeldzaam: loop de uitgangspunten en de verantwoording na voordat "
            "u daaruit een conclusie trekt."
        )
    else:
        onderdelen = [
            _aantal(sum(1 for punt in punten if punt.ernst == ernst), enkelvoud, meervoud)
            for ernst, enkelvoud, meervoud in (
                (KRITIEK, "kritiek punt", "kritieke punten"),
                (WAARSCHUWING, "waarschuwing", "waarschuwingen"),
                (SIGNAAL, "signaal", "signalen"),
            )
            if any(punt.ernst == ernst for punt in punten)
        ]
        inleiding = (
            f"Op basis van {bron} zijn "
            f"{_aantal(len(punten), 'aandachtspunt', 'aandachtspunten')} benoemd die "
            f"beoordeling vragen: {' en '.join(onderdelen)}."
        )
    staart = []
    if niet_mogelijk:
        staart.append(
            f"{_aantal(niet_mogelijk, 'controle kon', 'controles konden')} niet worden "
            "uitgevoerd; zie “Wat niet kon worden vastgesteld”"
        )
    if afgehandeld:
        werkwoord = "staat" if afgehandeld == 1 else "staan"
        staart.append(
            f"{_aantal(afgehandeld, 'bevinding is', 'bevindingen zijn')} al beoordeeld en "
            f"{werkwoord} achteraan"
        )
    if staart:
        inleiding = f"{inleiding} Daarnaast: {'; '.join(staart)}."

    slot = ""
    if len(punten) > SAMENVATTING_MAX:
        slot = (
            f"De overige {len(punten) - SAMENVATTING_MAX} punten staan hieronder, in dezelfde "
            "volgorde."
        )
    return Sectie(
        kop="Samenvatting",
        inleiding=inleiding,
        regels=tuple(punt.aanduiding for punt in punten[:SAMENVATTING_MAX]),
        genummerd=bool(punten),
        slot=slot,
    )


def _sectie_niet_mogelijk(punten: tuple[Punt, ...]) -> Sectie:
    if punten:
        inleiding = (
            "Deze controles konden op deze bestanden niet worden uitgevoerd. Dat is geen "
            "vaststelling dat er niets aan de hand is: waar de controle ontbreekt, ontbreekt "
            "ook het bewijs. Wilt u er wel een uitspraak over, dan is een export met de "
            "ontbrekende gegevens nodig."
        )
    else:
        inleiding = (
            "Elke controle van deze tool kon op deze bestanden worden uitgevoerd. Dat zegt "
            "niets over controles die deze tool niet kent; zie de verantwoording."
        )
    return Sectie(kop="Wat niet kon worden vastgesteld", inleiding=inleiding, punten=punten)


def _sectie_afgehandeld(punten: tuple[Punt, ...]) -> Sectie:
    return Sectie(
        kop="Al beoordeeld",
        inleiding=(
            "Deze bevindingen zijn van een beoordeling voorzien. Ze staan hier met de "
            "vastgelegde status en notitie, zodat in het dossier terug te vinden is wat er "
            "mee is gedaan."
        ),
        punten=punten,
    )


def _rgs_regel(huidig: Auditfile) -> str:
    rekeningen = len(huidig.accounts)
    if not rekeningen:
        return (
            "Het bestand levert geen rekeningschema, dus elke selectie op rekening berust op "
            "wat er in de boekingen staat."
        )
    met_rgs = int((huidig.accounts["RGScode"].astype(str).str.strip() != "").sum())
    if met_rgs == rekeningen:
        return (
            f"Alle {rekeningen} rekeningen hebben een RGS-code, dus de indeling berust op een "
            "code uit het bestand en niet op een omschrijving."
        )
    aandeel = met_rgs / rekeningen * 100
    return (
        f"{met_rgs} van de {rekeningen} rekeningen hebben een RGS-code ({aandeel:.0f}%); voor "
        "de rekeningen zonder code beslist de omschrijving. Bij elk punt staat welke methode "
        "is gebruikt."
    )


def _sectie_verantwoording(
    huidig: Auditfile, bevindingen: pd.DataFrame, materialiteit: Materialiteit
) -> Sectie:
    from .capability import NIVEAU_NAAM, openstaande_posten_niveau

    niveau, uitleg = openstaande_posten_niveau(huidig)
    naam = NIVEAU_NAAM.get(niveau, "onbekend")
    regels = [
        f"Openstaande posten zijn te bepalen op niveau {niveau}: {naam.lower()}. {uitleg}",
        _rgs_regel(huidig),
    ]
    if not bevindingen.empty:
        onder = int((~bevindingen["boven_drempel"].astype(bool)).sum())
        regels.append(
            f"{onder} van de {len(bevindingen)} bevindingen liggen onder de "
            f"materialiteitsdrempel van {euro(materialiteit.drempel)}. Ze zijn als zodanig "
            "gemarkeerd en niet weggelaten."
        )
        if "status" in bevindingen.columns:
            openstaand = int((bevindingen["status"] == TE_BEOORDELEN).sum())
            regels.append(
                f"{openstaand} van de {len(bevindingen)} bevindingen hebben nog geen "
                "vastgelegde beoordeling."
            )
    regels.append(
        "Dit memorandum is door de tool samengesteld uit het auditfile. Het vervangt geen "
        "dossiercontrole en is geen accountantsverklaring."
    )
    return Sectie(kop="Verantwoording", regels=tuple(regels))


# --- Opbouw -----------------------------------------------------------------


def bouw_memorandum(
    huidig: Auditfile,
    bevindingen: pd.DataFrame,
    materialiteit: Materialiteit,
    vorig: Auditfile | None = None,
    opsteller: str = "",
    opgesteld_op: date | None = None,
) -> Memorandum:
    """Zet de bevindingen om naar een leesbaar memorandum.

    ``bevindingen`` is de tabel van ``verzamel_bevindingen()``, bij voorkeur met
    de beoordeling erin via ``pas_review_toe()``; staat die kolom er niet, dan
    wordt zij hier toegevoegd zodat elk punt een status heeft.

    ``materialiteit`` is verplicht en niet optioneel: de drempel staat in het
    document, en die moet dezelfde zijn als waarmee de tabel is gemaakt. Een
    eigen standaardwaarde zou een drempel noemen waarop niet is gemeten.
    """
    if "status" not in bevindingen.columns:
        bevindingen = pas_review_toe(bevindingen)
    gesorteerd = op_gewicht(bevindingen)

    aandacht: dict[str, list[Punt]] = {KRITIEK: [], WAARSCHUWING: [], SIGNAAL: []}
    onmogelijk: list[Punt] = []
    afgehandeld: list[Punt] = []

    # De nummering loopt door over het hele document, zodat een punt met één
    # nummer aan te wijzen is.
    nummer = 0
    rijen = [rij for _, rij in gesorteerd.iterrows()]

    def _volgende(rij: pd.Series) -> Punt:
        nonlocal nummer
        nummer += 1
        return _naar_punt(rij, nummer)

    # Eerst de punten die beoordeling vragen, want die krijgen de laagste
    # nummers en de samenvatting verwijst ernaar.
    for rij in rijen:
        ernst = _tekst(rij.get("ernst"))
        status = _tekst(rij.get("status")) or TE_BEOORDELEN
        if ernst == NIET_MOGELIJK or status in AFGEHANDELDE_STATUSSEN:
            continue
        if ernst in aandacht:
            aandacht[ernst].append(_volgende(rij))
    for rij in rijen:
        if _tekst(rij.get("ernst")) == NIET_MOGELIJK:
            onmogelijk.append(_volgende(rij))
    for rij in rijen:
        ernst = _tekst(rij.get("ernst"))
        status = _tekst(rij.get("status")) or TE_BEOORDELEN
        if ernst != NIET_MOGELIJK and status in AFGEHANDELDE_STATUSSEN:
            afgehandeld.append(_volgende(rij))

    te_beoordelen = tuple(
        punt for ernst in (KRITIEK, WAARSCHUWING, SIGNAAL) for punt in aandacht[ernst]
    )

    secties: list[Sectie] = [
        _sectie_dossier(huidig, vorig, opsteller, opgesteld_op or date.today()),
        _sectie_uitgangspunten(materialiteit),
        _sectie_samenvatting(
            te_beoordelen, len(onmogelijk), len(afgehandeld), vorig is not None
        ),
    ]
    for ernst in (KRITIEK, WAARSCHUWING, SIGNAAL):
        if aandacht[ernst]:
            secties.append(Sectie(kop=ERNST_KOP[ernst], punten=tuple(aandacht[ernst])))
    secties.append(_sectie_niet_mogelijk(tuple(onmogelijk)))
    if afgehandeld:
        secties.append(_sectie_afgehandeld(tuple(afgehandeld)))
    secties.append(_sectie_verantwoording(huidig, bevindingen, materialiteit))

    jaar = huidig.boekjaar or "onbekend boekjaar"
    naam = huidig.bedrijfsnaam or "onderneming niet vermeld"
    return Memorandum(
        titel=f"Reviewmemorandum auditfile {jaar}",
        ondertitel=f"{naam}, opgesteld uit de auditfile-analyse",
        secties=tuple(secties),
    )


# --- Uitvoer ----------------------------------------------------------------


def _punt_regels(punt: Punt) -> list[str]:
    regels = [f"**{punt.nummer}. {punt.aanduiding}**", ""]
    if punt.toelichting:
        regels += [punt.toelichting, ""]
    if punt.herkomst:
        regels += [f"*{punt.herkomst}*", ""]
    if punt.beoordeling:
        regels += [punt.beoordeling, ""]
    return regels


def naar_markdown(memo: Memorandum) -> str:
    """Het memorandum als Markdown-tekst."""
    uit: list[str] = [f"# {memo.titel}", ""]
    if memo.ondertitel:
        uit += [memo.ondertitel, ""]
    for sectie in memo.secties:
        uit += [f"## {sectie.kop}", ""]
        if sectie.inleiding:
            uit += [sectie.inleiding, ""]
        if sectie.kenmerken:
            uit += [f"- **{label}:** {waarde}" for label, waarde in sectie.kenmerken]
            uit.append("")
        if sectie.regels:
            for index, regel in enumerate(sectie.regels, start=1):
                uit.append(f"{index}. {regel}" if sectie.genummerd else f"- {regel}")
            uit.append("")
        if sectie.slot:
            uit += [sectie.slot, ""]
        for punt in sectie.punten:
            uit += _punt_regels(punt)
    return "\n".join(uit).rstrip() + "\n"


def memorandum_markdown(
    huidig: Auditfile,
    bevindingen: pd.DataFrame,
    materialiteit: Materialiteit,
    vorig: Auditfile | None = None,
    opsteller: str = "",
    opgesteld_op: date | None = None,
) -> str:
    """De opbouw en de uitvoer in één stap, voor de app en de export."""
    return naar_markdown(
        bouw_memorandum(
            huidig,
            bevindingen,
            materialiteit,
            vorig=vorig,
            opsteller=opsteller,
            opgesteld_op=opgesteld_op,
        )
    )


# De stijlen uit de standaardsjabloon van python-docx, bij elkaar zodat de
# opmaak op een plek te wijzigen is en niet verspreid door de renderer.
DOCX_LIJST = "List Bullet"
DOCX_GENUMMERD = "List Paragraph"
DOCX_ONDERTITEL = "Subtitle"


def _docx_document():
    """Een leeg Word-document, met de afhankelijkheid pas hier geladen.

    python-docx staat in ``requirements.txt``, maar de import staat niet
    bovenaan deze module: een omgeving zonder die afhankelijkheid verliest dan
    alleen de Word-uitvoer en niet de hele app.
    """
    try:
        from docx import Document
    except ModuleNotFoundError as fout:  # pragma: no cover - hangt aan de omgeving
        raise ModuleNotFoundError(
            "De Word-uitvoer vraagt python-docx. Die staat in requirements.txt; "
            "installeer de afhankelijkheden met pip install -r requirements.txt."
        ) from fout
    return Document()


def _docx_cursief(document, tekst: str) -> None:
    document.add_paragraph().add_run(tekst).italic = True


def _docx_punt(document, punt: Punt) -> None:
    """Een punt in het Word-document.

    Een punt is hier een kop en niet een vetgedrukte regel zoals in Markdown.
    Dat is opmaak en geen formulering, en het levert een document op dat in het
    navigatievenster van Word te doorlopen is.
    """
    document.add_heading(f"{punt.nummer}. {punt.aanduiding}", level=3)
    if punt.toelichting:
        document.add_paragraph(punt.toelichting)
    if punt.herkomst:
        _docx_cursief(document, punt.herkomst)
    if punt.beoordeling:
        document.add_paragraph(punt.beoordeling)


def naar_docx(memo: Memorandum) -> bytes:
    """Het memorandum als Word-bestand.

    De tweede renderer op dezelfde ``Memorandum``: hier staat geen enkele zin,
    alleen de omzetting naar Word-opmaak. Wijzigt de formulering, dan wijzigt
    zij in ``bouw_memorandum()`` en komt zij hier vanzelf mee.
    """
    document = _docx_document()
    document.core_properties.title = memo.titel
    # De opsteller staat als kenmerk in het stuk zelf. In de
    # documenteigenschappen komt geen naam: die reist mee met het bestand en
    # zegt niets over wie het heeft beoordeeld.
    document.core_properties.author = "Auditfile Analyzer"

    document.add_heading(memo.titel, level=0)
    if memo.ondertitel:
        document.add_paragraph(memo.ondertitel, style=DOCX_ONDERTITEL)
    for sectie in memo.secties:
        document.add_heading(sectie.kop, level=1)
        if sectie.inleiding:
            document.add_paragraph(sectie.inleiding)
        for label, waarde in sectie.kenmerken:
            alinea = document.add_paragraph(style=DOCX_LIJST)
            alinea.add_run(f"{label}: ").bold = True
            alinea.add_run(waarde)
        for index, regel in enumerate(sectie.regels, start=1):
            if sectie.genummerd:
                # Het nummer staat in de tekst en niet in de automatische
                # nummering van Word: het verwijst naar het nummer van het punt
                # verderop, en dat verband mag niet verschuiven zodra iemand in
                # het document een regel toevoegt of weghaalt.
                document.add_paragraph(f"{index}. {regel}", style=DOCX_GENUMMERD)
            else:
                document.add_paragraph(regel, style=DOCX_LIJST)
        if sectie.slot:
            document.add_paragraph(sectie.slot)
        for punt in sectie.punten:
            _docx_punt(document, punt)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def memorandumnaam(huidig: Auditfile, extensie: str = "md") -> str:
    """Bestandsnaam voor de download, zonder klantnaam.

    Zoals bij de Excel-export: de naam van de onderneming hoort niet in een
    bestandsnaam die in een downloadmap terechtkomt.
    """
    jaar = "".join(teken for teken in str(huidig.boekjaar) if teken.isalnum()) or "boekjaar"
    return f"reviewmemorandum_{jaar}.{extensie.lstrip('.')}"
